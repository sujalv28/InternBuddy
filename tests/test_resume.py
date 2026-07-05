import io

import pytest
from docx import Document
from fpdf import FPDF

import resume


def test_extract_drive_id_file_form():
    link = "https://drive.google.com/file/d/1A2B3C_dEF/view?usp=sharing"
    assert resume.extract_drive_id(link) == "1A2B3C_dEF"


def test_extract_drive_id_open_form():
    link = "https://drive.google.com/open?id=XyZ-123"
    assert resume.extract_drive_id(link) == "XyZ-123"


def test_extract_drive_id_invalid():
    with pytest.raises(resume.ResumeError):
        resume.extract_drive_id("https://example.com/nope")


def test_extract_text_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Skilled in Python and machine learning")
    data = bytes(pdf.output())
    text = resume.extract_text(data)
    assert "Python" in text


def test_extract_text_docx():
    doc = Document()
    doc.add_paragraph("Experienced in data science and SQL")
    buf = io.BytesIO()
    doc.save(buf)
    text = resume.extract_text(buf.getvalue())
    assert "data science" in text


def test_get_resume_text_uses_download(monkeypatch):
    doc = Document()
    doc.add_paragraph("hello resume")
    buf = io.BytesIO()
    doc.save(buf)
    monkeypatch.setattr(resume, "download_resume", lambda link, timeout=30: buf.getvalue())
    assert "hello resume" in resume.get_resume_text("https://drive.google.com/file/d/ID/view")


class _Resp:
    def __init__(self, status_code=200, content=b"", ctype="", text=""):
        self.status_code = status_code
        self.content = content
        self.headers = {"Content-Type": ctype}
        self.text = text


class _FakeSession:
    """Serves queued responses keyed by a substring of the requested URL."""

    def __init__(self, routes):
        self.routes = routes
        self.headers = {}

    def get(self, url, timeout=30, params=None):
        for needle, resp in self.routes.items():
            if needle in url:
                return resp
        raise AssertionError(f"unexpected URL: {url}")


def test_download_resume_uploaded_pdf(monkeypatch):
    session = _FakeSession({
        "drive.google.com/uc": _Resp(content=b"%PDF-data", ctype="application/pdf"),
    })
    monkeypatch.setattr(resume.requests, "Session", lambda: session)
    assert resume.download_resume("https://drive.google.com/file/d/ID/view") == b"%PDF-data"


def test_download_resume_google_doc_falls_back_to_export(monkeypatch):
    session = _FakeSession({
        "drive.google.com/uc": _Resp(status_code=500, ctype="text/html"),
        "docs.google.com/document": _Resp(content=b"%PDF-exported", ctype="application/pdf"),
    })
    monkeypatch.setattr(resume.requests, "Session", lambda: session)
    assert resume.download_resume("https://drive.google.com/file/d/ID/view") == b"%PDF-exported"


def test_download_resume_all_endpoints_fail(monkeypatch):
    session = _FakeSession({
        "drive.google.com/uc": _Resp(status_code=500, ctype="text/html"),
        "docs.google.com/document": _Resp(status_code=404, ctype="text/html"),
    })
    monkeypatch.setattr(resume.requests, "Session", lambda: session)
    with pytest.raises(resume.ResumeError):
        resume.download_resume("https://drive.google.com/file/d/ID/view")
